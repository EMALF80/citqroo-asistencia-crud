from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_DIR = Path(__file__).resolve().parents[1]
OUT_PATH = PROJECT_DIR / "docs" / "Guia_Estudio_Proyecto_Final_CITQROO.docx"
REPO_URL = "https://github.com/EMALF80/citqroo-asistencia-crud"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table, color="D9E2EC"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(table):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)

    for margin_name, value in (("top", "100"), ("bottom", "100"), ("start", "120"), ("end", "120")):
        node = tbl_cell_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), value)
        node.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_cell_margins(table)
    set_table_borders(table)

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        set_cell_shading(cell, "F2F4F7")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(15, 23, 42)

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)

    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbers(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_code_block(doc, code, title=None):
    if title:
        paragraph = doc.add_paragraph()
        paragraph.add_run(title).bold = True

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_cell_margins(table)
    set_table_borders(table, color="CBD5E1")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F8FAFC")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)

    for line_number, line in enumerate(code.strip("\n").split("\n")):
        if line_number > 0:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(30, 41, 59)

    doc.add_paragraph()


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_cell_margins(table)
    set_table_borders(table, color="BFDBFE")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "EFF6FF")
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(f"{label}: ")
    run.bold = True
    run.font.color.rgb = RGBColor(29, 78, 216)
    paragraph.add_run(text)
    doc.add_paragraph()


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color, before, after in [
        ("Heading 1", 16, RGBColor(46, 116, 181), 16, 8),
        ("Heading 2", 13, RGBColor(46, 116, 181), 12, 6),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    return doc


def build_document():
    doc = setup_document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Guia de estudio profunda del proyecto final")
    run.font.name = "Calibri"
    run.font.size = Pt(23)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 23, 42)

    subtitle = doc.add_paragraph()
    subtitle.add_run("Proyecto: ").bold = True
    subtitle.add_run("Control de Asistencia CITQROO - CRUD con Vue.js y json-server")
    subtitle.paragraph_format.space_after = Pt(8)

    add_table(
        doc,
        ["Dato", "Valor"],
        [
            ("Carpeta", str(PROJECT_DIR)),
            ("Repositorio esperado", REPO_URL),
            ("API fake", "json-server en http://localhost:3001"),
            ("Frontend", "Vue.js 3 + Vite"),
            ("Objetivo de esta guia", "Estudiar y defender el codigo sin depender de internet"),
        ],
        widths=[1.7, 5.8],
    )

    add_callout(
        doc,
        "Idea central",
        "La app no es solo una pantalla: Vue muestra formularios y tablas, fetch consume la API, "
        "json-server guarda los datos en db.json, y un proceso calcula incidencias usando horarios y asistencias.",
    )

    doc.add_heading("1. Que se construyo", level=1)
    doc.add_paragraph(
        "Se construyo una aplicacion web para el control de asistencia del Centro de Innovacion CITQROO. "
        "La practica pide una aplicacion Vue.js que consuma una API interna, use al menos tres entidades/tablas, "
        "tenga CRUD, use componentes, binding, directivas y conceptos de Vue como props, emit, hooks, metodos y computed."
    )
    add_bullets(
        doc,
        [
            "Entidad principal 1: alumnos.",
            "Entidad principal 2: proyectos.",
            "Entidad principal 3: asistencias.",
            "Entidad principal 4: incidencias.",
            "Catalogos auxiliares: carreras, docentes, horarios, tiposIncidencia, tiposParticipacion y otros recursos.",
        ],
    )
    add_callout(
        doc,
        "Respuesta corta para el profesor",
        "Mi proyecto registra alumnos, proyectos y asistencias. Luego un modulo procesa la asistencia del dia, "
        "compara cada registro contra el horario programado y genera incidencias como retardo, falta, salida anticipada o fuera de horario.",
    )

    doc.add_heading("2. Como ejecutarlo desde cero", level=1)
    add_callout(doc, "Importante", "Se necesitan dos terminales: una para la API y otra para Vue.")
    add_numbers(
        doc,
        [
            "Abrir PowerShell o la terminal integrada de Visual Studio Code.",
            "Entrar a la carpeta del proyecto con cd.",
            "Instalar dependencias si es la primera vez.",
            "Levantar la API con npm run api.",
            "Abrir otra terminal y levantar Vue con npm run dev.",
            "Abrir la URL que da Vite, por ejemplo http://127.0.0.1:5173 o el puerto disponible.",
        ],
    )
    add_code_block(
        doc,
        'cd "C:\\Users\\Asus\\Documents\\Framework frontend\\citqroo-asistencia-crud"\n'
        "npm install\n"
        "npm run api",
        "Terminal 1: API",
    )
    add_code_block(
        doc,
        'cd "C:\\Users\\Asus\\Documents\\Framework frontend\\citqroo-asistencia-crud"\n'
        "npm run dev",
        "Terminal 2: frontend Vue",
    )
    add_code_block(doc, "npm run build\nnpm audit --audit-level=moderate", "Comandos para comprobar")

    doc.add_heading("3. Modelo mental del proyecto", level=1)
    doc.add_paragraph(
        "Piensa en tres capas: db.json contiene datos, json-server los publica como API, y Vue consume esa API con fetch."
    )
    add_table(
        doc,
        ["Capa", "Archivo o herramienta", "Responsabilidad"],
        [
            ("Datos", "db.json", "Guarda arrays como alumnos, proyectos, asistencias e incidencias."),
            ("API", "json-server", "Convierte cada array de db.json en endpoints REST."),
            ("Servicio", "src/services/api.js", "Centraliza fetch para GET, POST, PUT y DELETE."),
            ("Vista principal", "src/App.vue", "Controla navegacion y estadisticas."),
            ("Componentes", "src/components/*.vue", "Formularios, tablas, dashboard y proceso."),
            ("Logica", "src/utils/incidencias.js", "Calcula incidencias segun horarios y registros."),
        ],
        widths=[1.5, 2.4, 3.6],
    )

    doc.add_heading("4. Mapa de archivos", level=1)
    add_table(
        doc,
        ["Archivo", "Para que sirve"],
        [
            ("package.json", "Define nombre, scripts y dependencias."),
            ("db.json", "Base de datos fake usada por json-server."),
            ("index.html", "HTML base donde Vue monta la app en #app."),
            ("vite.config.js", "Configura Vite y el alias @."),
            ("src/main.js", "Crea la aplicacion Vue y monta App.vue."),
            ("src/App.vue", "Componente padre: navegacion, titulo, estadisticas y vistas."),
            ("src/services/api.js", "Funciones reutilizables para consumir la API."),
            ("src/utils/incidencias.js", "Reglas para generar incidencias."),
            ("src/components/EntityTable.vue", "Tabla reutilizable."),
            ("src/components/StudentManager.vue", "CRUD de alumnos."),
            ("src/components/ProjectManager.vue", "CRUD de proyectos."),
            ("src/components/AttendanceManager.vue", "CRUD de asistencia."),
            ("src/components/IncidenceProcessor.vue", "Procesa y muestra incidencias."),
            ("src/styles.css", "Estilos generales y responsive."),
        ],
        widths=[2.6, 4.9],
    )

    doc.add_heading("5. package.json explicado", level=1)
    doc.add_paragraph("package.json es el centro de comandos del proyecto.")
    add_code_block(
        doc,
        '"scripts": {\n'
        '  "dev": "vite --host 127.0.0.1",\n'
        '  "api": "json-server --watch db.json --port 3001",\n'
        '  "build": "vite build",\n'
        '  "preview": "vite preview --host 127.0.0.1"\n'
        "}",
        "Scripts importantes",
    )
    add_bullets(
        doc,
        [
            "npm run dev: levanta el servidor de desarrollo de Vue.",
            "npm run api: levanta json-server y crea la API REST desde db.json.",
            "npm run build: genera la version final en dist.",
            "npm run preview: revisa la version compilada.",
        ],
    )
    add_table(
        doc,
        ["Dependencia", "Tipo", "Explicacion"],
        [
            ("vue", "dependencies", "Framework frontend usado para componentes y reactividad."),
            ("vite", "devDependencies", "Servidor de desarrollo y herramienta de build."),
            ("@vitejs/plugin-vue", "devDependencies", "Permite que Vite entienda .vue."),
            ("json-server", "devDependencies", "Crea una API fake REST desde db.json."),
        ],
        widths=[2.0, 1.7, 3.8],
    )

    doc.add_heading("6. db.json y la API fake", level=1)
    doc.add_paragraph(
        "db.json funciona como una base de datos sencilla. Cada propiedad principal es una tabla o recurso. "
        "json-server convierte esas propiedades en endpoints."
    )
    add_table(
        doc,
        ["Recurso", "Endpoint", "Operacion"],
        [
            ("alumnos", "/alumnos", "CRUD de estudiantes."),
            ("proyectos", "/proyectos", "CRUD de proyectos."),
            ("asistencias", "/asistencias", "CRUD de entradas y salidas."),
            ("incidencias", "/incidencias", "Resultado del proceso automatico."),
            ("carreras", "/carreras", "Catalogo usado por alumnos."),
            ("docentes", "/docentes", "Catalogo usado por proyectos."),
            ("horarios", "/horarios", "Reglas por alumno y dia."),
            ("tiposIncidencia", "/tiposIncidencia", "Catalogo de retardo, falta, etc."),
        ],
        widths=[1.7, 2.1, 3.7],
    )
    add_callout(
        doc,
        "Clave para defender",
        "json-server no valida relaciones como una base SQL real. La relacion se hace en Vue buscando IDs, por ejemplo carreraId.",
    )

    doc.add_heading("7. Conceptos de Vue que usa el proyecto", level=1)
    add_table(
        doc,
        ["Concepto", "Donde aparece", "Explicacion"],
        [
            ("ref", "App.vue, managers", "Datos reactivos simples o arrays."),
            ("reactive", "StudentManager, ProjectManager, AttendanceManager", "Objetos de formulario."),
            ("computed", "App.vue, managers, IncidenceProcessor", "Datos derivados automaticos."),
            ("onMounted", "App.vue y managers", "Carga datos al aparecer el componente."),
            ("props", "AppHeader, EntityTable, DashboardView", "Datos del padre al hijo."),
            ("emit", "AppHeader, EntityTable, managers", "Eventos del hijo al padre."),
            ("v-model", "Inputs y selects", "Une campos con variables reactivas."),
            ("v-for", "Tablas, menus, selects", "Repite elementos por cada item."),
            ("v-if/v-else-if/v-else", "App.vue, EntityTable", "Muestra secciones segun estado."),
        ],
        widths=[1.5, 2.4, 3.6],
    )

    doc.add_heading("8. App.vue: el componente padre", level=1)
    doc.add_paragraph(
        "App.vue coordina la aplicacion. Decide que pantalla se muestra, carga conteos del resumen y escucha eventos de los hijos."
    )
    add_code_block(
        doc,
        'const views = [\n'
        '  { id: "dashboard", label: "Resumen" },\n'
        '  { id: "students", label: "Alumnos" },\n'
        '  { id: "projects", label: "Proyectos" },\n'
        '  { id: "attendance", label: "Asistencia" },\n'
        '  { id: "incidences", label: "Incidencias" }\n'
        "];",
        "Menu de vistas",
    )
    add_code_block(
        doc,
        "function navigate(viewId) {\n"
        "  activeView.value = viewId;\n"
        "  window.location.hash = viewId;\n"
        "}",
        "Cambio de pantalla",
    )
    add_code_block(
        doc,
        '<DashboardView v-if="activeView === \'dashboard\'" :loading="loadingStats" :stats="stats" />\n'
        '<StudentManager v-else-if="activeView === \'students\'" @changed="handleDataChanged" />',
        "Directivas y eventos",
    )
    add_bullets(
        doc,
        [
            "activeView indica que pantalla se ve.",
            "currentTitle es computed y cambia el titulo segun activeView.",
            "loadStats consulta la API para llenar tarjetas.",
            "@changed se ejecuta cuando un hijo modifica datos.",
            "window.location.hash permite abrir #students, #projects, #attendance o #incidences.",
        ],
    )

    doc.add_heading("9. api.js: como se consume la API", level=1)
    doc.add_paragraph("api.js evita repetir fetch en todos los componentes.")
    add_code_block(
        doc,
        'const API_URL = "http://localhost:3001";\n\n'
        "async function request(path, options = {}) {\n"
        "  const response = await fetch(`${API_URL}${path}`, {\n"
        '    headers: { "Content-Type": "application/json", ...(options.headers || {}) },\n'
        "    ...options\n"
        "  });\n\n"
        "  if (!response.ok) {\n"
        "    throw new Error(`Error ${response.status}: ${response.statusText}`);\n"
        "  }\n\n"
        "  return response.status === 204 ? null : response.json();\n"
        "}",
        "Funcion request",
    )
    add_table(
        doc,
        ["Funcion", "Metodo HTTP", "Ejemplo"],
        [
            ("getResource('alumnos')", "GET", "Trae todos los alumnos."),
            ("createResource('alumnos', payload)", "POST", "Crea un alumno."),
            ("updateResource('alumnos', id, payload)", "PUT", "Actualiza un alumno."),
            ("deleteResource('alumnos', id)", "DELETE", "Elimina un alumno."),
        ],
        widths=[2.5, 1.5, 3.5],
    )

    doc.add_heading("10. EntityTable.vue: tabla reutilizable", level=1)
    doc.add_paragraph(
        "EntityTable recibe columnas y filas por props. Cuando se presiona Editar o Eliminar, emite eventos al padre."
    )
    add_code_block(
        doc,
        "defineProps({\n"
        "  columns: { type: Array, required: true },\n"
        "  rows: { type: Array, required: true },\n"
        "  loading: { type: Boolean, default: false },\n"
        "  emptyMessage: { type: String, default: \"No hay registros para mostrar.\" }\n"
        "});\n\n"
        'const emit = defineEmits(["edit", "delete"]);',
        "Props y emit",
    )
    add_code_block(
        doc,
        '<tr v-for="row in rows" :key="row.id">\n'
        '  <td v-for="column in columns" :key="column.key">\n'
        "    {{ getCell(row, column) }}\n"
        "  </td>\n"
        '  <td class="row-actions">\n'
        '    <button @click="emit(\'edit\', row)">Editar</button>\n'
        '    <button @click="emit(\'delete\', row)">Eliminar</button>\n'
        "  </td>\n"
        "</tr>",
        "v-for y eventos",
    )

    doc.add_heading("11. Patron CRUD usado", level=1)
    doc.add_paragraph("Alumnos, proyectos y asistencia siguen el mismo patron.")
    add_numbers(
        doc,
        [
            "emptyForm define campos vacios.",
            "ref guarda datos, loading, editingId y message.",
            "reactive guarda el formulario.",
            "loadData trae datos desde json-server.",
            "submitForm crea o actualiza segun editingId.",
            "editRow llena el formulario con la fila.",
            "deleteRow elimina y recarga la tabla.",
            "emit('changed') avisa a App.vue.",
        ],
    )
    add_code_block(
        doc,
        "async function submitForm() {\n"
        "  const payload = {\n"
        "    nombreCompleto: form.nombreCompleto.trim(),\n"
        "    numeroControl: form.numeroControl.trim(),\n"
        "    carreraId: Number(form.carreraId),\n"
        "    semestre: Number(form.semestre)\n"
        "  };\n\n"
        "  if (editingId.value) {\n"
        "    await updateResource(\"alumnos\", editingId.value, { ...payload, id: editingId.value });\n"
        "  } else {\n"
        "    await createResource(\"alumnos\", payload);\n"
        "  }\n\n"
        "  resetForm();\n"
        "  await loadData();\n"
        "  emit(\"changed\");\n"
        "}",
        "Crear o actualizar alumno",
    )
    add_callout(doc, "Clave", "editingId es la diferencia entre crear y editar.")

    doc.add_heading("12. StudentManager.vue explicado", level=1)
    add_bullets(
        doc,
        [
            "Carga alumnos y carreras.",
            "El formulario usa nombreCompleto, numeroControl, carreraId y semestre.",
            "tableRows combina cada alumno con el nombre de su carrera.",
            "El select de carrera usa v-for.",
            "Al guardar convierte carreraId y semestre con Number(...).",
        ],
    )
    add_code_block(
        doc,
        "const tableRows = computed(() =>\n"
        "  alumnos.value.map((alumno) => ({\n"
        "    ...alumno,\n"
        "    carreraNombre:\n"
        "      carreras.value.find((carrera) => Number(carrera.id) === Number(alumno.carreraId))?.nombre ||\n"
        "      \"Sin carrera\"\n"
        "  }))\n"
        ");",
        "Mostrar nombre de carrera",
    )

    doc.add_heading("13. ProjectManager.vue explicado", level=1)
    add_bullets(
        doc,
        [
            "Carga proyectos y docentes.",
            "El formulario usa nombre, descripcion, fechaInicio, fechaTermino y docenteId.",
            "Si fechaTermino esta vacia, muestra En proceso.",
            "docenteNombre se obtiene buscando docenteId en docentes.",
        ],
    )
    add_code_block(
        doc,
        "docenteNombre:\n"
        "  docentes.value.find((docente) => Number(docente.id) === Number(proyecto.docenteId))\n"
        "    ?.nombreCompleto || \"Sin docente\",\n"
        "fechaTermino: proyecto.fechaTermino || \"En proceso\"",
        "Datos derivados de proyecto",
    )

    doc.add_heading("14. AttendanceManager.vue explicado", level=1)
    add_bullets(
        doc,
        [
            "Administra registros de entrada y salida.",
            "Cada registro tiene fecha, hora, alumnoId y tipoRegistro.",
            "La tabla se ordena del registro mas reciente al mas antiguo.",
            "El select de alumno se llena desde /alumnos.",
        ],
    )

    doc.add_heading("15. IncidenceProcessor.vue explicado", level=1)
    doc.add_paragraph(
        "Este modulo ejecuta un proceso. Al presionar Procesar dia, borra las incidencias anteriores de esa fecha, "
        "calcula nuevas incidencias y las guarda en la API."
    )
    add_code_block(
        doc,
        "async function processIncidences() {\n"
        "  processing.value = true;\n"
        "  message.value = \"\";\n\n"
        "  const previous = incidencias.value.filter((incidencia) => incidencia.fecha === fecha.value);\n"
        "  await Promise.all(previous.map((incidencia) => deleteResource(\"incidencias\", incidencia.id)));\n\n"
        "  const generated = buildIncidences({\n"
        "    fecha: fecha.value,\n"
        "    alumnos: alumnos.value,\n"
        "    horarios: horarios.value,\n"
        "    asistencias: asistencias.value,\n"
        "    tiposIncidencia: tiposIncidencia.value\n"
        "  });\n\n"
        "  await Promise.all(generated.map((incidencia) => createResource(\"incidencias\", incidencia)));\n"
        "  await loadData();\n"
        "  emit(\"changed\");\n"
        "}",
        "Proceso principal",
    )

    doc.add_heading("16. Reglas de incidencias", level=1)
    add_table(
        doc,
        ["Situacion", "Resultado"],
        [
            ("Hay registro pero no hay horario para ese dia", "Fuera de horario"),
            ("No hay entrada ni salida", "Falta"),
            ("No hay entrada", "Falta"),
            ("Entrada despues de mas de 20 minutos", "Falta"),
            ("Entrada despues de mas de 10 minutos", "Retardo"),
            ("No hay salida", "Falta"),
            ("Salida antes de mas de 15 minutos", "Falta"),
            ("Salida antes de mas de 5 minutos", "Salida anticipada"),
        ],
        widths=[4.0, 3.5],
    )
    add_code_block(
        doc,
        "if (delay > 20) {\n"
        "  // Falta\n"
        "} else if (delay > 10) {\n"
        "  // Retardo\n"
        "}",
        "Regla de entrada",
    )
    add_code_block(
        doc,
        "if (earlyDeparture > 15) {\n"
        "  // Falta\n"
        "} else if (earlyDeparture > 5) {\n"
        "  // Salida anticipada\n"
        "}",
        "Regla de salida",
    )

    doc.add_heading("17. Como modificar el proyecto", level=1)
    doc.add_heading("17.1 Agregar un campo nuevo a alumnos", level=2)
    add_numbers(
        doc,
        [
            "Abrir db.json y agregar el campo a cada alumno existente.",
            "Agregar el campo en emptyForm de StudentManager.vue.",
            "Agregarlo al payload de submitForm.",
            "Agregarlo en editRow.",
            "Agregar un input con v-model en el template.",
            "Agregar una columna en columns.",
        ],
    )
    add_code_block(
        doc,
        "const emptyForm = {\n"
        "  nombreCompleto: \"\",\n"
        "  numeroControl: \"\",\n"
        "  carreraId: \"\",\n"
        "  semestre: \"\",\n"
        "  telefono: \"\"\n"
        "};",
        "emptyForm con telefono",
    )
    add_code_block(
        doc,
        '<label>\n'
        "  Telefono\n"
        '  <input v-model="form.telefono" type="tel" />\n'
        "</label>",
        "Input nuevo",
    )

    doc.add_heading("17.2 Agregar un modulo CRUD nuevo", level=2)
    add_numbers(
        doc,
        [
            "Confirmar que db.json tenga el array del recurso.",
            "Crear un componente Manager copiando el patron de StudentManager.vue.",
            "Cambiar el recurso de API.",
            "Definir emptyForm, columns, loadData, submitForm, editRow y deleteRow.",
            "Importar el componente en App.vue.",
            "Agregar una vista en views.",
            "Agregar v-else-if para mostrarlo.",
        ],
    )
    add_code_block(
        doc,
        'import TeacherManager from "./components/TeacherManager.vue";\n\n'
        "const views = [\n"
        '  { id: "dashboard", label: "Resumen" },\n'
        '  { id: "teachers", label: "Docentes" }\n'
        "];\n\n"
        '<TeacherManager v-else-if="activeView === \'teachers\'" @changed="handleDataChanged" />',
        "Ejemplo para agregar docentes",
    )

    doc.add_heading("17.3 Cambiar el puerto de la API", level=2)
    add_numbers(
        doc,
        [
            "Cambiar package.json: json-server --watch db.json --port 3001.",
            "Cambiar src/services/api.js: API_URL.",
            "Reiniciar npm run api y npm run dev.",
        ],
    )
    add_code_block(
        doc,
        '// package.json\n"api": "json-server --watch db.json --port 3002"\n\n'
        "// src/services/api.js\n"
        'const API_URL = "http://localhost:3002";',
        "Ejemplo de cambio de puerto",
    )

    doc.add_heading("17.4 Cambiar una regla de incidencia", level=2)
    add_code_block(
        doc,
        "if (delay > 15) {\n"
        "  // Falta\n"
        "} else if (delay > 5) {\n"
        "  // Retardo\n"
        "}",
        "Regla modificada de ejemplo",
    )
    add_callout(doc, "Cuidado", "Despues de cambiar reglas hay que presionar Procesar dia otra vez.")

    doc.add_heading("18. Que hacer si el profe borra algo", level=1)
    add_table(
        doc,
        ["Si borra...", "Que revisar"],
        [
            ("Un input", "El form correspondiente y el label con v-model."),
            ("Una columna", "El array columns del componente correspondiente."),
            ("Boton editar/eliminar", "EntityTable.vue."),
            ("Una vista del menu", "views en App.vue y el bloque v-if/v-else-if."),
            ("Conexion API", "API_URL en src/services/api.js y npm run api."),
            ("Datos iniciales", "db.json."),
            ("Reglas de incidencias", "src/utils/incidencias.js."),
        ],
        widths=[2.5, 5.0],
    )

    doc.add_heading("19. Errores comunes y solucion", level=1)
    add_table(
        doc,
        ["Error", "Causa probable", "Solucion"],
        [
            ("Failed to fetch", "API apagada o puerto incorrecto.", "Ejecutar npm run api y revisar API_URL."),
            ("404 en /alumnos", "Recurso mal escrito.", "Revisar que db.json tenga alumnos."),
            ("Pantalla en blanco", "Error de Vue o import.", "Revisar consola del navegador."),
            ("npm no se reconoce", "Node no esta en PATH.", "Abrir terminal nueva o reinstalar Node.js."),
            ("Port already in use", "Puerto ocupado.", "Cerrar proceso anterior o cambiar puerto."),
            ("No se guardan cambios", "API apagada o fetch falla.", "Ver terminal de json-server."),
            ("Se duplican incidencias", "No se borraron antes.", "Revisar previous y deleteResource."),
        ],
        widths=[2.0, 2.7, 2.8],
    )

    doc.add_heading("20. Preguntas probables del profesor", level=1)
    qa = [
        ("Que hace tu proyecto?", "Administra alumnos, proyectos, asistencias e incidencias para CITQROO."),
        ("Donde esta la API?", "La levanta json-server leyendo db.json en http://localhost:3001."),
        ("Que metodo HTTP usas para crear?", "POST mediante createResource."),
        ("Que metodo HTTP usas para actualizar?", "PUT mediante updateResource."),
        ("Que metodo HTTP usas para eliminar?", "DELETE mediante deleteResource."),
        ("Por que usas computed?", "Para datos derivados como nombres de catalogos, filtros y ordenamientos."),
        ("Para que sirve emit?", "Para que un hijo avise al padre que ocurrio una accion."),
        ("Para que sirve onMounted?", "Para cargar datos cuando el componente aparece."),
        ("Donde esta el CRUD de alumnos?", "En src/components/StudentManager.vue."),
        ("Donde estan las reglas de incidencias?", "En src/utils/incidencias.js."),
        ("Que pasa si la API esta apagada?", "fetch falla y no se cargan datos."),
        ("Por que dos terminales?", "Una mantiene la API y otra mantiene el frontend."),
    ]
    for question, answer in qa:
        paragraph = doc.add_paragraph()
        paragraph.add_run("Pregunta: ").bold = True
        paragraph.add_run(question)
        paragraph = doc.add_paragraph()
        paragraph.add_run("Respuesta: ").bold = True
        paragraph.add_run(answer)

    doc.add_heading("21. Ejercicios de practica", level=1)
    add_numbers(
        doc,
        [
            "Ejecuta el proyecto desde cero con dos terminales.",
            "Agrega un alumno nuevo y confirma que aparece en Resumen.",
            "Edita el semestre de un alumno.",
            "Elimina un registro de asistencia y explica que metodo HTTP se uso.",
            "Procesa incidencias para 2026-06-05 y explica cada resultado.",
            "Agrega un campo telefono a alumnos.",
            "Cambia el umbral de retardo de 10 a 5 minutos.",
            "Explica EntityTable.vue: props, rows, columns y emit.",
        ],
    )

    doc.add_heading("22. Fragmentos clave para memorizar", level=1)
    add_code_block(
        doc,
        'import { createApp } from "vue";\n'
        'import App from "./App.vue";\n'
        'import "./styles.css";\n\n'
        'createApp(App).mount("#app");',
        "src/main.js",
    )
    add_code_block(
        doc,
        'getResource("alumnos");\n'
        'createResource("alumnos", payload);\n'
        'updateResource("alumnos", id, payload);\n'
        'deleteResource("alumnos", id);',
        "Servicio API",
    )
    add_code_block(
        doc,
        "const form = reactive({ ...emptyForm });\n"
        '<input v-model="form.nombreCompleto" type="text" required />',
        "Formulario reactivo",
    )
    add_code_block(
        doc,
        '<EntityTable\n'
        '  :columns="columns"\n'
        '  :rows="tableRows"\n'
        '  @edit="editRow"\n'
        '  @delete="deleteRow"\n'
        "/>",
        "Props y emit",
    )

    doc.add_heading("23. Guion corto para presentar", level=1)
    doc.add_paragraph(
        "Buenas, este es mi proyecto final. Elegi el proceso de control de asistencia del Centro de Innovacion CITQROO. "
        "La aplicacion esta hecha con Vue.js 3 y consume una API fake creada con json-server. La base se encuentra en db.json "
        "y contiene alumnos, proyectos, asistencias, horarios, catalogos e incidencias. La app tiene modulos CRUD para alumnos, "
        "proyectos y asistencia. Tambien tiene un modulo de incidencias que procesa una fecha, compara los registros con los horarios "
        "y genera incidencias segun las reglas del documento. Use componentes reutilizables, props, emit, v-model, v-for, v-if, "
        "computed, onMounted y metodos async con fetch."
    )

    doc.add_page_break()
    doc.add_heading("24. Anexo visual", level=1)
    shot_dir = PROJECT_DIR / "docs" / "screenshots"
    shots = [
        ("Resumen", "01-dashboard.png"),
        ("Alumnos", "02-alumnos.png"),
        ("Proyectos", "03-proyectos.png"),
        ("Asistencia", "03-asistencia.png"),
        ("Incidencias", "04-incidencias.png"),
    ]
    for title_text, filename in shots:
        path = shot_dir / filename
        doc.add_heading(title_text, level=2)
        if path.exists():
            doc.add_picture(str(path), width=Inches(6.8))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            doc.add_paragraph(f"Captura no encontrada: {filename}")

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build_document())
